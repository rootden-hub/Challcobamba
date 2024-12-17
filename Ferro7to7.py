import streamlit as st
from PIL import Image


from datetime import timedelta
from datetime import datetime


import matplotlib.patches as mpatches
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import re

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches
from docx.shared import Pt
from docx import Document
from io import BytesIO

# ****************** Extraer la fecha del nombre del archivo
def get_reporte_date(file_path):
    # Obtener solo el nombre del archivo
    file_name = file_path.split("/")[-1]
    # Buscar el patrón de fecha DD MM YY
    date_match = re.search(r'(\d{2}) (\d{2}) (\d{2})', file_name)
    
    if date_match:
        day = date_match.group(1)
        month = int(date_match.group(2))
        # Asumir que los años son del 2000 en adelante
        year = "20" + date_match.group(3)

        # Convertir el mes numérico a nombre en español
        months = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio", 
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
        ]
        month_name = months[month - 1]
        
        # Formatear la fecha
        formatted_date = f"{day} de {month_name} del {year}"
    else:
        # Si no se encuentra la fecha en el nombre del archivo
        formatted_date = None
    
    return formatted_date
    
def get_and_increment_date(file_path):
    # Obtener solo el nombre del archivo
    file_name = file_path.split("/")[-1]
    
    # Buscar el patrón de fecha en formato 'DD MM YY'
    date_match = re.search(r'(\d{2}) (\d{2}) (\d{2})', file_name)
    
    if date_match:
        # Extraer día, mes y año
        day = date_match.group(1)
        month = date_match.group(2)
        year = "20" + date_match.group(3)  # Suponiendo que el año es del 2000 en adelante
        
        # Crear la fecha en formato 'DD-MM-YYYY'
        date_str = f"{day}-{month}-{year}"
        
        # Convertir la fecha en un objeto datetime
        current_date = datetime.strptime(date_str, "%d-%m-%Y")
        
        # Sumar un día a la fecha
        incremented_date = current_date + timedelta(days=1)
        
        # Convertir la fecha incrementada nuevamente a formato string 'DD-MM-YYYY'
        incremented_date_str = incremented_date.strftime("%d-%m-%Y")
        
        return date_str, incremented_date_str
    else:
        return None, None    


# Función para formatear el tiempo como horas:minutos
def format_duration(td):
    if isinstance(td, pd.Timedelta):
        total_seconds = td.total_seconds()
        hours = int(total_seconds // 3600)
        minutes = int((total_seconds % 3600) // 60)
        return f"{hours:02d}:{minutes:02d}"
    return "00:00"  # En caso de que el valor no sea un Timedelta


def generate_daily_report(caution_df, alarm_df, report_date):
    # Verificar si los DataFrames están vacíos antes de continuar
    if caution_df.empty and alarm_df.empty:
        print("Ambos DataFrames están vacíos. No se generará el reporte.")
        return None  # Return None if both DataFrames are empty

    # Si ambos DataFrames no están vacíos, eliminar la primera y última fila solo del primero
    if not caution_df.empty and not alarm_df.empty:
        caution_df = caution_df.iloc[1:-1]

    # Convert 'Date' to datetime for easier manipulation solo si no están vacíos
    if not caution_df.empty:
        caution_df['Date'] = pd.to_datetime(caution_df['Date'])
    if not alarm_df.empty:
        alarm_df['Date'] = pd.to_datetime(alarm_df['Date'])

    # Concatenar los DataFrames
    combined_df = pd.concat([caution_df, alarm_df], ignore_index=True)
    
    # Ordenar las filas por la columna 'Date'
    combined_df = combined_df.sort_values(by='Date', ascending=True)
     # Convertir 'Date' a datetime si no está en ese formato (si es necesario)
    combined_df['Date'] = pd.to_datetime(combined_df['Date'])

    # Definir el rango de horas: entre 07:00 AM y 07:00 AM del siguiente día
    # Creamos una nueva columna que extrae solo la hora de 'Date'
    combined_df['Hour'] = combined_df['Date'].dt.hour + combined_df['Date'].dt.minute / 60

    # Filtrar solo las filas con hora entre 07:00 y 07:00 (pasando de las 07:00 AM de un día a las 07:00 AM del siguiente día)
    combined_df = combined_df[(combined_df['Hour'] >= 7) & (combined_df['Hour'] < 7 + 24)]
    # Ordenar las filas por la columna 'Date', y si hay fechas iguales, por 'Type' (Start primero)
    combined_df['Type_priority'] = combined_df['Type'].apply(lambda x: 0 if x == 'Start' else 1)
    combined_df = combined_df.sort_values(by=['Date', 'Type_priority'], ascending=[True, True])
    combined_df = combined_df.drop(columns=['Type_priority'])  # Eliminar columna auxiliar

  


    # Crear columna de 'Duration' en formato min:segundos
    durations = []
    for i in range(len(combined_df) - 1):
        end_time = combined_df.iloc[i + 1]['Date'] 
        start_time = combined_df.iloc[i]['Date']
        duration = end_time - start_time
        durations.append(duration)

    # Añadir una fila vacía para la última duración (no hay un siguiente evento)
    durations.append('')

    combined_df['Duration'] = durations




    # Actualizar la columna 'Status' según las nuevas reglas definidas
    def update_status(row):
        if row['Duration'] == pd.Timedelta(0):
            return 'Grey'

        if row['Type'] == '-':
            return 'Free-White'
        elif row['Description'] == 'Caution' and row['Type'] == 'Start':
            return 'Yellow'
        elif row['Description'] == 'Caution' and row['Type'] == 'End':
            if isinstance(row['Duration'], pd.Timedelta) and row['Duration'] < pd.Timedelta(hours=1):
                return 'Grey'
            elif isinstance(row['Duration'], pd.Timedelta) and row['Duration'] >= pd.Timedelta(hours=1):
                return 'Caution-White'
        elif row['Description'] == 'Alarm' and row['Type'] == 'Start':
            return 'Red'
        elif row['Description'] == 'Alarm' and row['Type'] == 'End':
            return 'Yellow'
        return ''

    # Aplicar la función para actualizar la columna 'Status'
    combined_df['Status'] = combined_df.apply(update_status, axis=1)

    # Variables para las duraciones totales por tipo
    total_yellow_duration = pd.Timedelta(0)
    total_grey_duration = pd.Timedelta(0)
    total_red_duration = pd.Timedelta(0)
    total_white_duration = pd.Timedelta(0)

    # Sumar las duraciones de cada tipo
    for _, row in combined_df.iterrows():
        if isinstance(row['Duration'], pd.Timedelta):  # Ignorar duraciones vacías
            if row['Status'] == 'Yellow':
                total_yellow_duration += row['Duration']
            elif row['Status'] == 'Grey':
                total_grey_duration += row['Duration']
            elif row['Status'] == 'Red':
                total_red_duration += row['Duration']
            elif row['Status'] in ['Free-White', 'Caution-White']:
                total_white_duration += row['Duration']

    status_colors = {
        'Grey': 'grey',
        'Free-White': 'white',
        'Yellow': 'yellow',
        'Caution-White': 'white',
        'Red': 'red'
    }

    fig, ax = plt.subplots(figsize=(15, 8))

    # Adjust x-axis limits to the specific day
    #agregamos+ pd.Timedelta(hours=7)------------------------------------------------------------------------------------------------------
    start_date = combined_df['Date'].min().normalize() + pd.Timedelta(hours=7)
    end_date = start_date + pd.Timedelta(days=1)
    ax.set_xlim(start_date, end_date)

    # Set x-axis ticks every 2 hours
    ax.xaxis.set_major_locator(mdates.HourLocator(interval=1))
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%H:%M'))

    # Remove x-axis label
    ax.set_xlabel('')

    # Set y-axis label
    ax.set_ylabel('Tipo de Alerta')

    # Plot the bars
    for i, row in combined_df.iterrows():
        # Condición para omitir duraciones de cero
        if isinstance(row['Duration'], pd.Timedelta) and row['Duration'] > pd.Timedelta(0):
            start = row['Date']
            duration_in_minutes = row['Duration'].total_seconds() / 60
            color = status_colors.get(row['Status'], 'black')
            ax.barh(0, width=duration_in_minutes, left=start, color=color, edgecolor='none')
            
            # Formato de duración como hh:mm
            total_seconds = row['Duration'].total_seconds()
            hours = int(total_seconds // 3600)
            minutes = int((total_seconds % 3600) // 60)
            duration_text = f"{hours:02d}:{minutes:02d}"
            
            # Añadir texto verticalmente centrado o hacia abajo para barras blancas
            if color in ['white']:  # Verifica si la barra es blanca
                ax.text(start + row['Duration'] / 2, -0.3, duration_text, ha='center',
                        fontsize=9, color='black', rotation=90)
            else:
                ax.text(start + row['Duration'] / 2, 0, duration_text, ha='center',
                        fontsize=9, color='black', rotation=90)
           
                
    ax.set_yticks([])
    ax.set_title(f'{report_date} - Sensores Ferrobamba', fontsize=16, pad=20, loc='left')

    # Filtrar solo las alertas de interés para la leyenda
    legend_status = ['Red', 'Yellow', 'Grey']
    legend_patches = [
        mpatches.Patch(color=color, label='Alerta Roja' if status == 'Red' else 
                                        'Alerta Amarilla' if status == 'Yellow' else 
                                        'Libre entre Alertas <=1hr' if status == 'Grey' else status) 
        for status, color in status_colors.items() if status in legend_status
    ]

    # Crear la leyenda arriba a la derecha, encima de la gráfica
    ax.legend(handles=legend_patches, loc='lower right', bbox_to_anchor=(1, 1.05), ncol=5)

    # Crear la tabla con las métricas
    total_alert_duration = total_yellow_duration + total_red_duration
    total_combined_duration = total_alert_duration + total_grey_duration


    table_data = [
        ['Tiempo Alerta Amarilla', format_duration(total_yellow_duration)],
        ['Tiempo Alerta Roja', format_duration(total_red_duration)],
        ['Total Alertas (Amarilla+Roja)', format_duration(total_alert_duration)],
        ['Total Tiempo Libre Entre Alertas (<=1Hr)', format_duration(total_grey_duration)],
        ['Total Alertas + Tiempo Libre Entre Alertas', format_duration(total_combined_duration)]
    ]


    # Añadir la tabla a la gráfica
    table = ax.table(cellText=table_data, loc='bottom', cellLoc='center', colLoc='center', bbox=[0.1, -0.533, 0.8, 0.3])

    # Personalizar las celdas de la tabla
    for (i, j), cell in table.get_celld().items():
        if j == -1:  # Títulos de las filas
            cell.set_fontsize(10)
            cell.set_text_props(weight='bold')
            cell.set_facecolor('#ffcccb')
            cell.set_text_props(color='black')
        if i == 0:  # Títulos de las columnas
            cell.set_fontsize(10)
            cell.set_text_props(weight='bold')
            cell.set_facecolor('#4CAF50')
            cell.set_text_props(color='white')

    plt.tight_layout()

    return ax


def plot_eventos(df, report_date):
    # Conversión de fechas
    df['Start'] = pd.to_datetime(df['Start'])
    df['hora'] = df['Start'].dt.hour

    # Contar los eventos por hora y por tipo
    contador_tipo_1 = df[df['Description'] == 'Caution Flash'].groupby('hora').size()  # Amarillo
    contador_tipo_2 = df[df['Description'] == 'Alarm Flash'].groupby('hora').size()  # Roja
    contador_tipo_3 = df[df['Description'] == 'Warning Flash'].groupby('hora').size()  # Naranja

    # Sumar Alarm Flash y Caution Flash
    contador_tipo_2_y_3 = contador_tipo_2.add(contador_tipo_3, fill_value=0)

    # Crear un DataFrame con ambos conteos
    #Se agega para modficar linea de tiempo-------------------------------------------------------------------------------------------------
    horas = list(range(7, 24)) + list(range(0, 7))  # 7 AM a 7 AM
    conteos = pd.DataFrame({
        'Amarilla': contador_tipo_1.reindex(horas, fill_value=0),
        'Roja': contador_tipo_2_y_3.reindex(horas, fill_value=0)
    #----------------------------------------------------------------------------------------------------------
    #conteos = pd.DataFrame({
     #   'Amarilla': contador_tipo_1.reindex(range(24), fill_value=0),
      #  'Roja': contador_tipo_2_y_3.reindex(range(24), fill_value=0)
    }).fillna(0)
    
    # Calcular el total, promedio y máximo para cada tipo de evento
    total_tipo_1 = conteos['Amarilla'].sum()
    promedio_tipo_1 = conteos['Amarilla'].mean()
    maximo_tipo_1 = conteos['Amarilla'].max()

    total_tipo_2_y_3 = conteos['Roja'].sum()
    promedio_tipo_2_y_3 = conteos['Roja'].mean()
    maximo_tipo_2_y_3 = conteos['Roja'].max()

    # Definir la posición de las barras
    x = np.arange(len(conteos))  # Las posiciones de las horas
    width = 0.35  # Ancho de las barras

    # Crear el subplot
    fig, ax = plt.subplots(figsize=(15, 8))

    # Plot de las dos series de datos
    bars1 = ax.bar(x - width/2, conteos['Amarilla'], width, label='Amarilla', color='yellow')
    bars2 = ax.bar(x + width/2, conteos['Roja'], width, label='Roja', color='red')
#----------------------------------------------------------------------
    

    initial_date, incremented_date = get_and_increment_date(file_path)
   # Etiquetas y título
    ax.set_xlabel('Horas del día')
    ax.set_ylabel('Eventos')
    ax.set_title(f'Frecuencia de descargas eléctricas por hora del día {initial_date}\nSensores Ferrobamba', fontsize=16, pad=20)
    ax.set_xticks(x)
    ax.set_xticklabels([f'{h:02d}:00' for h in range(24)])


    # Rotar las etiquetas del eje X
    plt.xticks(rotation=90)

    # Mostrar la cantidad de eventos encima de cada barra
    for bar in bars1:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, height, str(int(height)), ha='center', va='bottom', fontsize=10)

    for bar in bars2:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, height, str(int(height)), ha='center', va='bottom', fontsize=10)

    # Leyenda
    ax.legend(title='Tipo de Evento', loc='upper right', fontsize=9, title_fontsize=11, shadow=True, fancybox=True, facecolor='silver', edgecolor='black')

    # Agregar la grilla
    ax.grid(True, which='both', axis='both', linestyle='--', color='gray', alpha=0.5)

    # Crear la tabla con métricas
    table_data = [
        [total_tipo_1, total_tipo_2_y_3],
        [round(conteos['Amarilla'].mean()), round(conteos['Roja'].mean())],
        [maximo_tipo_1, maximo_tipo_2_y_3]
    ]
    row_labels = ['Total', 'Promedio', 'Máximo']
    column_labels = ['Alerta Amarilla', 'Alerta Roja']

    # Añadir la tabla a la gráfica
    table = ax.table(cellText=table_data, rowLabels=row_labels, colLabels=column_labels, loc='bottom', cellLoc='center', colLoc='center', bbox=[0.1, -0.533, 0.8, 0.3])

    # Personalizar las celdas de la tabla
    for (i, j), cell in table.get_celld().items():
        if j == -1:  # Títulos de las filas
            cell.set_fontsize(10)
            cell.set_text_props(weight='bold')
            cell.set_facecolor('#ffcccb')
            cell.set_text_props(color='black')
        if i == 0:  # Títulos de las columnas
            cell.set_fontsize(10)
            cell.set_text_props(weight='bold')
            cell.set_facecolor('#4CAF50')
            cell.set_text_props(color='white')

    plt.tight_layout()

    # Devolver el gráfico sin mostrarlo
    return ax


def generate_report(df, file_name):
    report_date = get_reporte_date(file_name)

    # Verificar las descripciones únicas en la columna 'Description'
    unique_descriptions = df['Description'].dropna().unique()  # Ignorar valores nulos

    # Crear un diccionario para almacenar DataFrames por descripción
    description_dfs = {}

    for description in unique_descriptions:
        # Filtrar el DataFrame original para cada descripción
        filtered_df = df[df['Description'] == description]

        # Crear un nuevo DataFrame con pares Start-End si 'End' no tiene valores nulos
        if filtered_df['End'].isna().any():
            # Si hay valores nulos en 'End', agregar el DataFrame original sin reorganizar
            description_dfs[description] = filtered_df
        else:
            # Reorganizar el DataFrame si no hay valores nulos en 'End'
            reorganized_data = []
            for _, row in filtered_df.iterrows():
                # Agregar el par Start
                if not pd.isna(row['Start']):
                    reorganized_data.append({
                        'Date': row['Start'],
                        'Description': description,
                        'Type': 'Start',
                        'Status': 'Active'
                    })
                # Agregar el par End
                reorganized_data.append({
                    'Date': row['End'],
                    'Description': description,
                    'Type': 'End',
                    'Status': 'Active'
                })

            # Convertir los datos reorganizados a un DataFrame
            reorganized_df = pd.DataFrame(reorganized_data)

            # Obtener la fecha del primer registro y establecer la hora 00:00
            first_date = pd.to_datetime(reorganized_df['Date'].iloc[0]).normalize()+ pd.Timedelta(hours=7)

            start_of_day = first_date.strftime('%m/%d/%y %H:%M %p')

            # Crear un registro al inicio
            reorganized_df = pd.concat([
                pd.DataFrame([{
                    'Date': start_of_day,
                    'Description': description,
                    'Type': '-',
                    'Status': 'Active'
                }]),
                reorganized_df
            ], ignore_index=True)

            # Obtener la fecha del último registro y sumar un día para establecer la hora 00:00
            last_date = pd.to_datetime(reorganized_df['Date'].iloc[-1]).normalize() + timedelta(days=1)+ pd.Timedelta(hours=7)

            start_of_next_day = last_date.strftime('%m/%d/%y %H:%M %p')

            # Crear un registro al final
            reorganized_df = pd.concat([
                reorganized_df,
                pd.DataFrame([{
                    'Date': start_of_next_day,
                    'Description': description,
                    'Type': '-',
                    'Status': 'Active'
                }])
            ], ignore_index=True)

            # Convertir la columna 'Date' a datetime para asegurarse que se pueda calcular la duración
            reorganized_df['Date'] = pd.to_datetime(reorganized_df['Date'])

            # Crear columna de Duration en formato min:segundos
            durations = []
            for i in range(len(reorganized_df) - 1):
                end_time = reorganized_df.iloc[i + 1]['Date']
                duration = end_time - reorganized_df.iloc[i]['Date']
                durations.append(duration)

            # El último registro no tiene duración, así que lo dejamos vacío
            durations.append("")

            # Asignar la columna de duración al DataFrame reorganizado
            reorganized_df['Duration'] = durations

            # Agregar al diccionario
            description_dfs[description] = reorganized_df

    # Generar la gráfica Diaria
    caution_df = description_dfs.get('Caution', pd.DataFrame())  # Si no existe, devuelve un DataFrame vacío
    alarm_df = description_dfs.get('Alarm', pd.DataFrame())  # Si no existe, devuelve un DataFrame vacío
    ax_1 = generate_daily_report(caution_df, alarm_df, report_date)  # Get ax object from the report function
    ax_2 = plot_eventos(df, report_date)

    # Mostrar la primera gráfica en Streamlit
    st.subheader("Gráfica de Alertas por Día")
    st.pyplot(ax_1.figure)

    # Mostrar la segunda gráfica en Streamlit
    st.subheader("Frecuencia de Alertas por Hora")
    st.pyplot(ax_2.figure)


    # ******************************************

    img_buf_1 = BytesIO()
    ax_1.figure.savefig(img_buf_1, format='png')
    img_buf_1.seek(0)

    img_buf_2 = BytesIO()
    ax_2.figure.savefig(img_buf_2, format='png')
    img_buf_2.seek(0)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    header = doc.sections[0].header
    header_table = header.add_table(rows=1, cols=2, width=doc.sections[0].page_width)
    header_table.columns[0].width = Pt(50)

    cell_logo = header_table.cell(0, 0)
    cell_logo.paragraphs[0].add_run().add_picture("logo_doc.PNG", width=Pt(100))

    cell_title = header_table.cell(0, 1)
    header_table.cell(0, 1).width = Pt(1250)
    title_paragraph = cell_title.paragraphs[0]
    title_paragraph.add_run("REPORTE DIARIO DE ALERTAS POR DESCARGAS ELÉCTRICAS ATMOSFÉRICAS\n").bold = True
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle_paragraph = cell_title.add_paragraph(
        f"De: {report_date} 07:00 horas\tA: {report_date} 07:00 horas"
    )
    subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #doc.add_paragraph("Gráfica de alertas diarias:")
    doc.add_picture(img_buf_1, width=Pt(650))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #doc.add_paragraph("Gráfica de eventos por hora:")
    doc.add_picture(img_buf_2, width=Pt(650))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # Obtener la fecha actual para el footer
    today_date = datetime.now().strftime("%d de %B del %Y")

    # Configurar el pie de página
    footer = section.footer  # Acceder al footer de la sección
    footer_paragraph = footer.paragraphs[0]  # Crear un párrafo dentro del footer

    # Configurar el contenido del footer
    footer_paragraph.text = (
        f"{today_date} Supervisión de Mantenimiento Eléctrico-Telecomunicaciones"
    )
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar el texto

    # Ajustar el formato del texto del footer
    for run in footer_paragraph.runs:
        run.font.size = Pt(10)  # Ajustar el tamaño de la fuente

    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    st.download_button(
        label="Descargar Informe",
        data=doc_buffer,
        file_name=f"informe_generado-{report_date}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    

# Función para cargar y mostrar archivo
def cargar_archivo():
    archivo = st.file_uploader("Selecciona un archivo de texto (.txt)", type=["txt"])
    
    if archivo is not None:
        try:
            # Leer el archivo como DataFrame
            df = pd.read_csv(archivo, sep='\t')
            st.success(f"Archivo cargado correctamente: {archivo.name}")
            
            # Mostrar los datos en un DataFrame interactivo
            st.dataframe(df)

            
            # Botón para generar reportes (aunque no está implementado, lo mostramos)
            if st.button("Generar Reportes"):
                st.write("Generando reportes...")
                generate_report(df, archivo.name)
        
        except Exception as e:
            st.error(f"Ocurrió un error al leer el archivo: {e}")
    else:
        st.warning("Por favor, selecciona un archivo de texto (.txt)")

# Mostrar el logo (asegúrate de tener el archivo en la ruta correcta)
try:
    imagen = Image.open("logo.png")
    imagen = imagen.resize((100, 50), Image.Resampling.LANCZOS)
    st.image(imagen, use_container_width=False)
except FileNotFoundError:
    st.warning("No se encontró la imagen en la ruta especificada.")

# Título largo de la app
st.title("Generador de Reportes Diarios FERROBAMBA")


# Llamar a la función para cargar el archivo
cargar_archivo()
